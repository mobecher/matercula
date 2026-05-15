"""Test fixtures generated on-the-fly to keep the repo tiny."""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

FIXTURES = Path(__file__).parent / "fixtures"


def _build_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    # Page 1
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, height - 80, "Einfuehrung in Social Engineering")
    c.setFont("Helvetica", 12)
    c.drawString(
        72,
        height - 120,
        "Phishing bezeichnet den Versuch, an sensible Daten zu gelangen.",
    )
    c.drawString(
        72,
        height - 140,
        "Angreifer geben sich als vertrauenswuerdige Stelle aus.",
    )
    c.showPage()
    # Page 2
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, height - 80, "Schutzmassnahmen")
    c.setFont("Helvetica", 12)
    c.drawString(72, height - 120, "Mehrfaktor-Authentifizierung erhoeht die Sicherheit.")
    c.drawString(72, height - 140, "Nutzer sollten Absenderadressen pruefen.")
    c.showPage()
    c.save()


def _build_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Einleitung", level=1)
    doc.add_paragraph(
        "Dieser Text beschreibt grundlegende Konzepte digitaler Sicherheit."
    )
    doc.add_heading("Passwoerter", level=1)
    doc.add_paragraph("Lange Passphrasen sind sicherer als kurze Passwoerter.")
    doc.save(str(path))


@pytest.fixture(scope="session")
def sample_pdf_bytes() -> bytes:
    FIXTURES.mkdir(exist_ok=True)
    pdf = FIXTURES / "sample.pdf"
    if not pdf.exists():
        _build_pdf(pdf)
    return pdf.read_bytes()


@pytest.fixture(scope="session")
def sample_docx_bytes() -> bytes:
    FIXTURES.mkdir(exist_ok=True)
    docx = FIXTURES / "sample.docx"
    if not docx.exists():
        _build_docx(docx)
    return docx.read_bytes()


@pytest.fixture(scope="session")
def corrupt_pdf_bytes() -> bytes:
    return b"%PDF-1.4\nthis is not a real pdf\n%%EOF"


@pytest.fixture()
def client():
    from fastapi.testclient import TestClient

    from app.main import app

    with TestClient(app) as c:
        yield c


# Silence reportlab/pillow imports if io unused.
_ = io
